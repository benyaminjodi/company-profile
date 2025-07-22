import os
import re
import google.generativeai as genai
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
from datetime import datetime
from dotenv import load_dotenv
import streamlit as st

# Load API Key
load_dotenv()
# Gunakan secrets dari Streamlit Cloud jika ada, fallback ke .env
GOOGLE_API_KEY = st.secrets.get("GOOGLE_API_KEY") or os.getenv("GOOGLE_API_KEY")

# Konfigurasi Gemini API
genai.configure(api_key=GOOGLE_API_KEY)

# Struktur Bab II
document_structure = [
    {"title": "2.1 Latar Belakang dan Sejarah Perusahaan"},
    {"title": "2.2 Profil & Struktur Perusahaan", "subsections": [
        {"title": "2.2.1 Profil Perusahaan"},
        {"title": "2.2.2 Struktur Perusahaan"},
    ]},
    {"title": "2.3 Visi, Misi dan Tujuan Perusahaan", "subsections": [
        {"title": "2.3.1 Visi Perusahaan"},
        {"title": "2.3.2 Misi Perusahaan"},
        {"title": "2.3.3 Nilai Perusahaan"},
        {"title": "2.3.4 Tujuan Perusahaan"},
    ]},
    {"title": "2.4 Prinsip – Prinsip Tata Kelola Korporasi (Good Corporate Governance)"},
    {"title": "2.5 Evaluasi Pelaksanaan RJP Sebelumnya (2021–2025)", "subsections": [
        {"title": "2.5.1 Pencapaian Tujuan Perusahaan pada RJP 2021–2025"},
        {"title": "2.5.2 Pelaksanaan strategi dan kebijakan yang telah ditetapkan"},
    ]},
    {"title": "2.6 Benchmark Perusahaan Global Terminal Operator (GTO)"}
]

def upload_to_google_drive(local_file_path, drive_folder_id=None):
    gauth = GoogleAuth()
    gauth.LocalWebserverAuth()
    drive = GoogleDrive(gauth)

    file_drive = drive.CreateFile({'title': os.path.basename(local_file_path)})

    if drive_folder_id:
        file_drive['parents'] = [{"id": drive_folder_id}]

    file_drive.SetContentFile(local_file_path)
    file_drive.Upload()

    return file_drive['alternateLink']

def fetch_section_content(company_name, section_title, temperature=0.7, model_name="models/gemini-1.5-flash-latest"):
    prompt = f"""
Tulis bagian "{section_title}" untuk profil perusahaan "{company_name}".
Gunakan bahasa formal dan komprehensif. Tambahkan daftar sumber referensi jika relevan di bagian akhir sebagai list bernomor.
"""
    try:
        model = genai.GenerativeModel(model_name=model_name)
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"[GAGAL MENGAMBIL KONTEN: {e}]"

def split_content_and_sources(text):
    match = re.split(r"Sumber:|Referensi:", text, maxsplit=1, flags=re.IGNORECASE)
    content = match[0].strip()
    sources = match[1].strip().splitlines() if len(match) > 1 else []
    return content, sources

def add_table_of_contents(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldCharType')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)

def generate_full_company_profile_docx(company_name, upload=True, drive_folder_id=None, temperature=0.7, model_name="models/gemini-1.5-flash-latest"):
    doc = Document()

    doc.add_heading("BAB II  PENDAHULUAN", level=1)
    toc_paragraph = doc.add_paragraph()
    add_table_of_contents(toc_paragraph)
    doc.add_paragraph("(Klik kanan pada daftar isi di Microsoft Word > Update field untuk menampilkan halaman)")

    all_sources = []

    for section in document_structure:
        doc.add_heading(section["title"], level=2)
        if "subsections" in section:
            for sub in section["subsections"]:
                doc.add_heading(sub["title"], level=3)
                content = fetch_section_content(company_name, sub["title"], temperature, model_name)
                main_text, sources = split_content_and_sources(content)
                doc.add_paragraph(main_text)
                all_sources.extend(sources)
        else:
            content = fetch_section_content(company_name, section["title"], temperature, model_name)
            main_text, sources = split_content_and_sources(content)
            doc.add_paragraph(main_text)
            all_sources.extend(sources)

    if all_sources:
        doc.add_page_break()
        doc.add_heading("Daftar Sumber Referensi", level=1)
        for src in all_sources:
            doc.add_paragraph(src.strip(), style='List Number')

    safe_name = re.sub(r'[^\w\s-]', '', company_name).replace(" ", "_")
    filename = f"Bab_II_Profil_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)

    drive_url = None
    if upload:
        try:
            drive_url = upload_to_google_drive(filename, drive_folder_id)
            print(f"✅ File berhasil diupload ke Google Drive: {drive_url}")
        except Exception as e:
            print(f"❌ Gagal upload ke Google Drive: {e}")

    return filename, drive_url
